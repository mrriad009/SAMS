#!/usr/bin/env python3
"""Generate CSE 4100 Field Work report in NUBTK guideline format."""

from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

from generate_report_diagrams import generate_all as generate_diagrams

ASSETS_DIR = Path(__file__).parent / "report_assets"

# --- Student & supervisor details ---
STUDENT_NAME = "Md Mahamudul Islam Riad"
STUDENT_ID = "11220320898"
STUDENT_SECTION = "8E"
SUPERVISOR_NAME = "Arjan Ghosh"
SUPERVISOR_DESIGNATION = "Lecturer"
SUBMISSION_DATE = "June 2026"
TRAINING_LOCATION = (
    "Department of Computer Science and Engineering, "
    "Northern University of Business and Technology Khulna"
)
TRAINING_PERIOD = "March 2026 to June 2026"
REPORT_TITLE = (
    "Development of a Student Attendance Management System "
    "for University CSE Department"
)
OUTPUT_PATH = (
    "/Users/captainx/Downloads/SAMS/"
    "CSE4100_Field_Work_Report_Md_Mahamudul_Islam_Riad.docx"
)


def set_document_defaults(doc):
    section = doc.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)
    section.orientation = WD_ORIENT.PORTRAIT
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(1.25)

    style = doc.styles["Normal"]
    font = style.font
    font.name = "Times New Roman"
    font.size = Pt(12)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    pf = style.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    pf.space_after = Pt(6)


def styled_run(paragraph, text, *, bold=False, size=12):
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    return run


def add_paragraph(doc, text, *, bold=False, size=12, align=WD_ALIGN_PARAGRAPH.JUSTIFY, indent=0):
    p = doc.add_paragraph()
    p.alignment = align
    if indent:
        p.paragraph_format.left_indent = Inches(indent)
    styled_run(p, text, bold=bold, size=size)
    return p


def add_blank_lines(doc, count=1):
    for _ in range(count):
        doc.add_paragraph()


def add_page_break(doc):
    doc.add_page_break()


def add_centered_title(doc, text, size=14, bold=True, blanks_before=0, blanks_after=2):
    add_blank_lines(doc, blanks_before)
    add_paragraph(doc, text, bold=bold, size=size, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_blank_lines(doc, blanks_after)


def add_chapter(doc, number, title):
    add_page_break(doc)
    add_centered_title(doc, f"CHAPTER {number}", size=14, bold=True, blanks_before=0, blanks_after=2)
    add_centered_title(doc, title, size=14, bold=True, blanks_before=0, blanks_after=2)


def add_section(doc, text, level=1):
    indent = 0.25 if level == 2 else 0.5 if level == 3 else 0
    add_paragraph(doc, text, bold=True, align=WD_ALIGN_PARAGRAPH.JUSTIFY, indent=indent)


def add_body(doc, text):
    add_paragraph(doc, text)


def add_bullet_list(doc, items):
    for item in items:
        add_body(doc, f"• {item}")


def add_figure(doc, image_path, caption, width=Inches(5.8)):
    if not Path(image_path).exists():
        add_paragraph(doc, caption, align=WD_ALIGN_PARAGRAPH.CENTER)
        return

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(image_path), width=width)
    add_paragraph(doc, caption, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_blank_lines(doc, 1)


def add_table(doc, headers, rows, caption=None):
    if caption:
        add_paragraph(doc, caption, align=WD_ALIGN_PARAGRAPH.CENTER)
        add_blank_lines(doc, 1)

    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    for col, header in enumerate(headers):
        cell = table.rows[0].cells[col]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        styled_run(p, header, bold=True)

    for row_idx, row_data in enumerate(rows, start=1):
        for col_idx, value in enumerate(row_data):
            cell = table.rows[row_idx].cells[col_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            styled_run(p, value)

    add_blank_lines(doc, 1)


def build_report():
    generate_diagrams()
    doc = Document()
    set_document_defaults(doc)

    # ==================== TITLE PAGE ====================
    add_centered_title(doc, REPORT_TITLE, size=14, blanks_before=4, blanks_after=3)
    add_paragraph(doc, "by", align=WD_ALIGN_PARAGRAPH.CENTER)
    add_blank_lines(doc, 1)
    add_paragraph(doc, STUDENT_NAME, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_paragraph(doc, f"ID: {STUDENT_ID}", align=WD_ALIGN_PARAGRAPH.CENTER)
    add_blank_lines(doc, 2)
    add_paragraph(doc, "CSE 4100", align=WD_ALIGN_PARAGRAPH.CENTER)
    add_paragraph(doc, "Field Work / Industrial Training", align=WD_ALIGN_PARAGRAPH.CENTER)
    add_blank_lines(doc, 2)
    add_paragraph(doc, "Department of Computer Science and Engineering", align=WD_ALIGN_PARAGRAPH.CENTER)
    add_paragraph(doc, "Northern University of Business and Technology Khulna", align=WD_ALIGN_PARAGRAPH.CENTER)
    add_paragraph(doc, "Khulna 9100, Bangladesh", align=WD_ALIGN_PARAGRAPH.CENTER)
    add_blank_lines(doc, 1)
    add_paragraph(doc, SUBMISSION_DATE, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)

    # ==================== DECLARATION ====================
    add_page_break(doc)
    add_centered_title(doc, "Declaration", blanks_before=0, blanks_after=2)
    add_body(
        doc,
        f'This is to certify that the Field Work / Industrial Training work entitled '
        f'"{REPORT_TITLE}" has been carried out by {STUDENT_NAME} in the Department of '
        "Computer Science and Engineering, Northern University of Business and Technology "
        "Khulna, Khulna 9100, Bangladesh. The above Field Work / Industrial Training work "
        "or any part of this work has not been submitted anywhere for the award of any "
        "degree or diploma.",
    )
    add_blank_lines(doc, 2)
    add_paragraph(doc, "Signature of the Supervisor", align=WD_ALIGN_PARAGRAPH.LEFT)
    add_paragraph(doc, "Signature of the Student", align=WD_ALIGN_PARAGRAPH.LEFT)
    add_blank_lines(doc, 1)
    add_paragraph(doc, SUPERVISOR_NAME, align=WD_ALIGN_PARAGRAPH.LEFT)
    add_paragraph(doc, SUPERVISOR_DESIGNATION, align=WD_ALIGN_PARAGRAPH.LEFT)
    add_paragraph(doc, "Department of Computer Science and Engineering", align=WD_ALIGN_PARAGRAPH.LEFT)
    add_paragraph(doc, "Northern University of Business and Technology Khulna", align=WD_ALIGN_PARAGRAPH.LEFT)
    add_paragraph(doc, "Khulna 9100, Bangladesh", align=WD_ALIGN_PARAGRAPH.LEFT)
    add_blank_lines(doc, 1)
    add_paragraph(doc, STUDENT_NAME, align=WD_ALIGN_PARAGRAPH.LEFT)
    add_paragraph(doc, STUDENT_ID, align=WD_ALIGN_PARAGRAPH.LEFT)
    add_paragraph(doc, STUDENT_SECTION, align=WD_ALIGN_PARAGRAPH.LEFT)

    # ==================== ACKNOWLEDGEMENT ====================
    add_page_break(doc)
    add_centered_title(doc, "Acknowledgement", blanks_before=0, blanks_after=2)
    add_body(
        doc,
        "First of all, I would like to express my sincere gratitude to Almighty Allah for "
        "giving me the strength, patience, and opportunity to complete this field work and "
        "prepare this report.",
    )
    add_body(
        doc,
        f"I am especially grateful to my supervisor, {SUPERVISOR_NAME}, {SUPERVISOR_DESIGNATION}, "
        "Department of Computer Science and Engineering, Northern University of Business and "
        "Technology Khulna. From the early planning stage to the final review, he guided me "
        "with practical suggestions and honest feedback. His direction helped me stay focused "
        "and improve both the quality of the software and the clarity of this report.",
    )
    add_body(
        doc,
        "I would also like to thank the faculty members and staff of the CSE department for "
        "sharing their experience about classroom management. Their comments about attendance "
        "collection, section handling, and exam eligibility rules helped me understand the "
        "real problems behind this project.",
    )
    add_body(
        doc,
        "I am thankful to my classmates and class representatives of Section 8E, who gave me "
        "feedback during testing. Their day-to-day experience with manual attendance tracking "
        "helped me check whether the system was actually useful, not just technically complete.",
    )
    add_body(
        doc,
        "Finally, I owe thanks to my family and friends for their continuous encouragement "
        "throughout this training period.",
    )
    add_blank_lines(doc, 2)
    add_paragraph(doc, STUDENT_NAME, align=WD_ALIGN_PARAGRAPH.RIGHT)
    add_paragraph(doc, SUBMISSION_DATE, align=WD_ALIGN_PARAGRAPH.RIGHT)

    # ==================== CONTENTS ====================
    add_page_break(doc)
    add_centered_title(doc, "Contents", blanks_before=0, blanks_after=2)
    contents = [
        ("Title Page", "i"),
        ("Declaration", "ii"),
        ("Acknowledgement", "iii"),
        ("Contents", "iv"),
        ("List of Tables", "v"),
        ("List of Figures", "vi"),
        ("List of Abbreviations", "vii"),
        ("CHAPTER I  Introduction", "1"),
        ("    1.1  Introduction", "1"),
        ("    1.2  Motivation", "2"),
        ("    1.3  Objectives and Specific Aims", "3"),
        ("    1.4  Organization of the Report", "4"),
        ("CHAPTER II  Procedure / Methodology", "5"),
        ("    2.1  Requirements Analysis", "5"),
        ("    2.2  System Architecture and Design", "7"),
        ("    2.3  Technology Stack", "8"),
        ("    2.4  Database Design", "10"),
        ("    2.5  Development and Implementation Process", "11"),
        ("CHAPTER III  Results and Discussion", "14"),
        ("    3.1  Implemented System Modules", "14"),
        ("    3.2  Role-Based Features and Workflows", "16"),
        ("    3.3  Testing and Validation", "18"),
        ("    3.4  Challenges and Limitations", "19"),
        ("CHAPTER IV  Conclusions and Recommendations", "21"),
        ("    4.1  Conclusion", "21"),
        ("    4.2  Learning Outcomes", "22"),
        ("    4.3  Recommendations", "23"),
        ("References", "25"),
    ]
    for left, right in contents:
        p = doc.add_paragraph()
        p.paragraph_format.tab_stops.add_tab_stop(Inches(6.0))
        run_l = p.add_run(left)
        run_l.font.name = "Times New Roman"
        run_l.font.size = Pt(12)
        run_l.bold = left.startswith("CHAPTER")
        run_r = p.add_run("\t" + right)
        run_r.font.name = "Times New Roman"
        run_r.font.size = Pt(12)

    # ==================== LIST OF TABLES ====================
    add_page_break(doc)
    add_centered_title(doc, "List of Tables", blanks_before=0, blanks_after=2)
    add_table(
        doc,
        ["Table No.", "Description", "Page"],
        [
            ("2.1", "Technology Stack Summary", "8"),
            ("2.2", "User Roles and Permission Matrix", "9"),
            ("3.1", "Core System Modules and Description", "14"),
            ("3.2", "Attendance Status Types", "17"),
        ],
    )

    # ==================== LIST OF FIGURES ====================
    add_page_break(doc)
    add_centered_title(doc, "List of Figures", blanks_before=0, blanks_after=2)
    add_table(
        doc,
        ["Figure No.", "Description", "Page"],
        [
            ("2.1", "Three-Tier System Architecture", "7"),
            ("2.2", "Authentication and Session Flow", "11"),
            ("2.3", "Database Entity Relationship Diagram", "10"),
            ("3.1", "Attendance Marking Workflow", "16"),
            ("3.2", "Student Portal Dashboard Overview", "18"),
        ],
    )

    # ==================== LIST OF ABBREVIATIONS ====================
    add_page_break(doc)
    add_centered_title(doc, "List of Abbreviations", blanks_before=0, blanks_after=2)
    abbrevs = [
        ("API", "Application Programming Interface"),
        ("CR", "Class Representative"),
        ("CSE", "Computer Science and Engineering"),
        ("CRUD", "Create, Read, Update, Delete"),
        ("JWT", "JSON Web Token"),
        ("NUBTK", "Northern University of Business and Technology Khulna"),
        ("ORM", "Object-Relational Mapping"),
        ("REST", "Representational State Transfer"),
        ("UI", "User Interface"),
        ("UX", "User Experience"),
    ]
    for short, full in abbrevs:
        p = doc.add_paragraph()
        styled_run(p, f"{short}\t\t{full}")

    # ==================== CHAPTER I ====================
    add_chapter(doc, "I", "Introduction")

    add_section(doc, "1.1 Introduction")
    add_body(
        doc,
        "Class attendance is one of the basic academic requirements in almost every university. "
        "It affects exam eligibility, academic monitoring, and student discipline. Yet in many "
        "CSE departments, attendance is still recorded through paper registers, Excel files, "
        "or informal WhatsApp updates. These methods may work for a small group, but they "
        "become messy when the department grows across multiple semesters, sections, and courses.",
    )
    add_body(
        doc,
        "During my CSE 4100 field work, I developed a full-stack Student Attendance Management "
        "System (SAMS) for the Computer Science and Engineering department. The system supports "
        "daily attendance marking, course enrollment, class routines, announcements, reports, "
        "and public student lookup. It also provides separate access levels for administrators, "
        "class representatives (CRs), teachers, and students. The project was designed with "
        "real department-level use in mind, especially for academic operations at Northern "
        "University of Business and Technology Khulna (NUBTK).",
    )
    add_body(
        doc,
        "This report describes the full journey of the project — from identifying the problem "
        "to planning the system, building the modules, testing the features, and reviewing "
        "the final outcome. My aim is not only to explain what was built, but also to show "
        "why each part was needed and how it can help improve attendance management in a "
        "real university setting.",
    )

    add_section(doc, "1.2 Motivation")
    add_body(
        doc,
        "The idea for this project came from problems I saw around me every day. Students often "
        "struggle to know their exact attendance percentage before midterms or finals. CRs spend "
        "the first ten minutes of class collecting attendance on paper. Teachers sometimes keep "
        "separate records in notebooks, which makes verification difficult later.",
    )
    add_body(
        doc,
        "Another issue is transparency. When attendance data is scattered across different "
        "files and people, students cannot quickly check their own status. This creates confusion "
        "and unnecessary pressure near exam time. A single digital platform can reduce this "
        "problem by storing records in one place and showing the right information to the right user.",
    )
    add_body(
        doc,
        "I also wanted to build something that matches how a CSE department actually works. "
        "It is not enough to mark only present or absent. The system must handle semester-wise "
        "sections, course sessions, weekly routines, CR permissions, reports, and student "
        "self-service features. This field work gave me the chance to move beyond a classroom "
        "demo and work on a solution closer to real institutional needs.",
    )

    add_section(doc, "1.3 Objectives and Specific Aims")
    add_body(
        doc,
        "The main objective of this field work was to design and develop a reliable web-based "
        "attendance management system for a university CSE department. The system should reduce "
        "manual work, improve record accuracy, and give both staff and students a clear view "
        "of attendance status.",
    )

    add_section(doc, "1.3.1 Specific Aims", level=2)
    add_bullet_list(
        doc,
        [
            "To study the existing attendance workflow in a CSE department and identify its weaknesses.",
            "To design a role-based system for admin, teacher/CR, student, and public users.",
            "To build a secure authentication system using JWT access tokens and HTTP-only refresh cookies.",
            "To implement session-based attendance with present, absent, late, and excused statuses.",
            "To develop student, course, routine, announcement, and reporting modules in one platform.",
            "To support NUBTK academic structures such as semesters, sections, faculty data, and CR accounts.",
            "To test the system in a development environment and judge its readiness for academic use.",
        ],
    )

    add_section(doc, "1.4 Organization of the Report")
    add_body(
        doc,
        "This report is divided into four chapters. Chapter I introduces the background, motivation, "
        "and objectives. Chapter II explains the methodology, system design, technology selection, "
        "and implementation process. Chapter III presents the developed features, workflows, "
        "testing observations, and limitations. Chapter IV concludes the work, summarizes my "
        "learning outcomes, and suggests future improvements.",
    )

    # ==================== CHAPTER II ====================
    add_chapter(doc, "II", "Procedure / Methodology")

    add_section(doc, "2.1 Requirements Analysis")
    add_body(
        doc,
        "The first step was to understand how attendance is handled in practice. I talked with "
        "CRs, reviewed common classroom habits, and observed how faculty members expect attendance "
        "data to be stored and checked. From this study, the following requirements were identified:",
    )
    add_bullet_list(
        doc,
        [
            "Staff must be able to mark attendance quickly during or right after class.",
            "CR accounts should only access their own semester and section.",
            "Students should be able to view attendance history and percentage by course.",
            "Admins should manage students, courses, routines, reports, and system settings.",
            "Attendance must be stored course-wise and session-wise.",
            "The system should support announcements and routine viewing.",
            "A public lookup feature should allow basic attendance checking by student ID.",
        ],
    )
    add_body(
        doc,
        "After collecting these points, I grouped them into functional requirements such as "
        "attendance, enrollment, reporting, and announcements, and non-functional requirements "
        "such as security, usability, and maintainability. This step gave me a clear roadmap "
        "before starting development.",
    )

    add_section(doc, "2.2 System Architecture and Design")
    add_body(
        doc,
        "SAMS follows a three-tier architecture with a React frontend, an Express backend, and "
        "a PostgreSQL database hosted on NeonDB. The frontend handles user interaction, the backend "
        "handles business rules and security, and the database stores academic and attendance data.",
    )
    add_figure(
        doc,
        ASSETS_DIR / "fig_2_1_architecture.png",
        "Fig. 2.1: Three-Tier System Architecture",
    )
    add_body(
        doc,
        "On the frontend, TanStack Query manages server state, while Zustand stores authentication "
        "state. On the backend, the code is organized into routes, controllers, services, and "
        "middleware. This separation made debugging easier and allowed me to update one layer "
        "without breaking the whole system.",
    )
    add_body(
        doc,
        "The application also supports two UI modes. General mode is attendance-focused and useful "
        "for quick daily work by CRs. Advanced mode provides the full admin and teacher dashboard "
        "with courses, routines, settings, and analytics.",
    )

    add_section(doc, "2.3 Technology Stack")
    add_body(
        doc,
        "I selected the technologies based on development speed, community support, scalability, "
        "and suitability for a full-stack academic project. Table 2.1 summarizes the main stack.",
    )
    add_table(
        doc,
        ["Layer", "Technologies Used"],
        [
            ("Frontend", "React 19, TypeScript, Vite, Tailwind CSS, shadcn/ui"),
            ("State Management", "Zustand, TanStack Query"),
            ("Backend", "Node.js, Express 5, TypeScript"),
            ("Database", "PostgreSQL (NeonDB) with Drizzle ORM"),
            ("Authentication", "JWT access token + HTTP-only refresh cookie"),
            ("Security", "Helmet, CORS, bcrypt, express-rate-limit, Zod validation"),
            ("Optional Services", "Cloudinary (avatars), Resend (password reset email)"),
        ],
        caption="Table 2.1: Technology Stack Summary",
    )

    add_table(
        doc,
        ["User Role", "Main Permissions"],
        [
            ("Admin", "Full access to students, courses, attendance on any date, reports, and settings"),
            ("Teacher / CR", "Scoped staff panel; attendance only for today within assigned section"),
            ("Student", "Profile, attendance view, course enrollment, routine, and announcements"),
            ("Public", "Student lookup by roll number without login"),
        ],
        caption="Table 2.2: User Roles and Permission Matrix",
    )

    add_section(doc, "2.4 Database Design")
    add_body(
        doc,
        "The database schema was designed to reflect real academic relationships. Main tables include "
        "users, students, teachers, courses, student_courses, class_sessions, attendance, announcements, "
        "class_routine, and system_settings. Supporting tables such as refresh_tokens, password_reset_tokens, "
        "notifications, academic_faculty, and section_representatives were also added.",
    )
    add_body(
        doc,
        "The attendance table keeps one record per student per class session. This unique rule "
        "prevents duplicate entries and keeps the data consistent. Enrollment is handled through "
        "student_courses, so only enrolled students appear on a session attendance sheet.",
    )
    add_figure(
        doc,
        ASSETS_DIR / "fig_2_3_database_er.png",
        "Fig. 2.3: Database Entity Relationship Diagram",
        width=Inches(6.2),
    )

    add_section(doc, "2.5 Development and Implementation Process")
    add_body(
        doc,
        "I followed a step-by-step development approach. First, I built the database schema and "
        "authentication flow. Then I implemented core modules such as student management, course "
        "management, and attendance marking. After that, I added announcements, routine, reports, "
        "and public lookup.",
    )
    add_figure(
        doc,
        ASSETS_DIR / "fig_2_2_auth_flow.png",
        "Fig. 2.2: Authentication and Session Flow",
    )
    add_body(
        doc,
        "During implementation, I applied server-side validation and role checks on every sensitive "
        "route. For example, a CR cannot change system settings, delete students, or mark attendance "
        "for another section. These rules were enforced in middleware and service layers, not only "
        "in the user interface.",
    )
    add_body(
        doc,
        "To make testing realistic, NUBTK-specific data such as course catalogs, class routines, "
        "faculty records, section representatives, and CR accounts were imported through seed "
        "scripts from parsed markdown sources. This was much better than using dummy placeholder data.",
    )
    add_body(
        doc,
        f"The field work was conducted at {TRAINING_LOCATION} during {TRAINING_PERIOD}. "
        "My daily work included requirement review, feature implementation, debugging, UI "
        "improvement, testing with sample academic data, and preparing this report.",
    )

    # ==================== CHAPTER III ====================
    add_chapter(doc, "III", "Results and Discussion")

    add_section(doc, "3.1 Implemented System Modules")
    add_body(
        doc,
        "By the end of the training period, I completed a working full-stack attendance management "
        "platform. Table 3.1 lists the major modules and their purpose.",
    )
    add_table(
        doc,
        ["Module", "Description"],
        [
            ("Authentication", "Login, logout, token refresh, password reset, student registration"),
            ("Attendance", "Session-based attendance sheets with calendar navigation"),
            ("Student Portal", "Profile management, attendance history, course enrollment"),
            ("Course Management", "Course catalog, teacher assignment, enrollment handling"),
            ("Routine", "Weekly class schedule viewing and management"),
            ("Announcements", "Targeted notices for all, department, or section"),
            ("Reports", "Attendance trends, defaulter lists, course-wise statistics"),
            ("Public Lookup", "Student search and public attendance summary"),
            ("Settings", "Attendance threshold, academic year, semester, app mode"),
        ],
        caption="Table 3.1: Core System Modules and Description",
    )

    add_section(doc, "3.2 Role-Based Features and Workflows")
    add_body(
        doc,
        "One of the strongest parts of the system is the attendance marking workflow. Staff users "
        "select semester, section, course, and date, then mark each enrolled student as present, "
        "absent, late, or excused. In general mode, the attendance page is simplified for fast "
        "daily use with a calendar strip and quick toggles.",
    )
    add_figure(
        doc,
        ASSETS_DIR / "fig_3_1_attendance_workflow.png",
        "Fig. 3.1: Attendance Marking Workflow",
    )
    add_table(
        doc,
        ["Status", "Meaning"],
        [
            ("Present", "Student attended the class"),
            ("Absent", "Student did not attend the class"),
            ("Late", "Student attended after the allowed time"),
            ("Excused", "Student was absent with valid approval"),
        ],
        caption="Table 3.2: Attendance Status Types",
    )
    add_body(
        doc,
        "The student portal gives a dashboard with attendance summary, enrolled courses, routine, "
        "and announcements. This allows students to check their academic standing without asking "
        "the CR again and again.",
    )
    add_figure(
        doc,
        ASSETS_DIR / "fig_3_2_student_portal.png",
        "Fig. 3.2: Student Portal Dashboard Overview",
    )
    add_body(
        doc,
        "The reporting module lets staff filter attendance by course, section, semester, and date "
        "range. It can show course-wise percentages, 30-day trend data, and students below the "
        "attendance threshold. The default threshold is 75%, but it can be changed from settings.",
    )

    add_section(doc, "3.3 Testing and Validation")
    add_body(
        doc,
        "I carried out manual functional testing in the development environment. The main areas "
        "tested were login and role redirection, attendance submission, CR scope restriction, "
        "student enrollment, announcement targeting, and public student lookup.",
    )
    add_body(
        doc,
        "The test results were encouraging. The system correctly blocked a CR from marking attendance "
        "for another section, stopped teachers from submitting past-date attendance, and calculated "
        "attendance percentages from session records. Seed data for NUBTK courses, routines, and "
        "Section 8E students made the testing process much closer to real use.",
    )
    add_body(
        doc,
        "Formal unit testing and full production deployment were not completed within this field "
        "work period. Even so, the current version is stable enough for controlled pilot use. "
        "More testing is still needed before department-wide adoption.",
    )

    add_section(doc, "3.4 Challenges and Limitations")
    add_body(
        doc,
        "I faced several challenges during development. Designing role permissions that work for "
        "both admin and CR users without confusing the interface took extra time. Importing large "
        "academic datasets from markdown sources into structured database tables also required "
        "careful checking.",
    )
    add_body(
        doc,
        "Some limitations remain. The system still depends on manual attendance marking instead "
        "of biometric or RFID integration. Password reset through email needs external service "
        "setup. Full mobile optimization and offline support were not finished in this phase. "
        "Complete production deployment with full user onboarding was also outside the immediate "
        "scope of this field work.",
    )

    # ==================== CHAPTER IV ====================
    add_chapter(doc, "IV", "Conclusions and Recommendations")

    add_section(doc, "4.1 Conclusion")
    add_body(
        doc,
        "This field work resulted in a practical Student Attendance Management System for a "
        "university CSE department. The project addressed common problems of manual attendance "
        "tracking by providing a centralized, role-based, and web-accessible solution.",
    )
    add_body(
        doc,
        "The final system supports attendance marking, student and course management, routine "
        "scheduling, announcements, reporting, and public lookup. It is especially suitable for "
        "departments with multiple semesters and sections, such as the CSE department at NUBTK. "
        "Overall, the work showed me that a well-planned digital attendance platform can save time, "
        "improve transparency, and support better academic monitoring.",
    )

    add_section(doc, "4.2 Learning Outcomes")
    add_body(
        doc,
        "This training gave me valuable hands-on experience in full-stack web development. I learned "
        "how to turn a real institutional problem into software requirements and then into working modules.",
    )
    add_bullet_list(
        doc,
        [
            "Practical experience in React, TypeScript, Express, and PostgreSQL development.",
            "Better understanding of authentication, authorization, and secure API design.",
            "Experience in database modeling for academic systems.",
            "Knowledge of role-based access control in real applications.",
            "Improved skill in debugging, code organization, and feature-based development.",
            "A clearer sense of how software should match actual academic workflow.",
        ],
    )

    add_section(doc, "4.3 Recommendations")
    add_body(
        doc,
        "Based on the work completed and the challenges observed, I suggest the following steps "
        "for future development:",
    )
    add_bullet_list(
        doc,
        [
            "Deploy the system on production hosting and run a pilot with selected sections.",
            "Add biometric, QR code, or RFID-based attendance capture to reduce manual input.",
            "Develop a dedicated mobile application for CRs and students.",
            "Introduce automated email or SMS alerts for low attendance and important announcements.",
            "Add export options for PDF and Excel attendance reports.",
            "Conduct formal usability testing with faculty members and students before full rollout.",
            "Prepare user manuals and short training sessions for admin and CR accounts.",
        ],
    )

    # ==================== REFERENCES ====================
    add_page_break(doc)
    add_centered_title(doc, "References", size=12, bold=True, blanks_before=0, blanks_after=2)
    references = [
        'R. Fielding, "Architectural Styles and the Design of Network-based Software Architectures," Ph.D. dissertation, Univ. of California, Irvine, 2000.',
        'M. Fowler, Patterns of Enterprise Application Architecture. Boston, MA, USA: Addison-Wesley, 2002.',
        'I. Sommerville, Software Engineering, 10th ed. Boston, MA, USA: Pearson, 2016.',
        'D. Flanagan, JavaScript: The Definitive Guide, 7th ed. Sebastopol, CA, USA: O\'Reilly Media, 2020.',
        'PostgreSQL Global Development Group, "PostgreSQL Documentation," 2025. [Online]. Available: https://www.postgresql.org/docs/',
        'Drizzle Team, "Drizzle ORM Documentation," 2025. [Online]. Available: https://orm.drizzle.team/docs/overview',
        'Meta Open Source, "React Documentation," 2025. [Online]. Available: https://react.dev/',
        'OpenJS Foundation, "Express.js Guide," 2025. [Online]. Available: https://expressjs.com/',
    ]
    for ref in references:
        add_body(doc, ref)

    doc.save(OUTPUT_PATH)
    print(f"Report saved to: {OUTPUT_PATH}")


if __name__ == "__main__":
    build_report()
